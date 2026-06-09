"""
booklet_generator.py — place in repo root alongside main.py
Generates a .docx audit working paper booklet for an engagement.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _heading(doc, text, level=1, color="003366"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16)
        )
    return p


def _kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(str(value) if value is not None else "—")
    r2.font.size = Pt(10)


def _tbl_header(table, headers, bg="003366", fg="FFFFFF"):
    """Apply bold coloured header row to a table."""
    row = table.rows[0]
    for cell, header in zip(row.cells, headers):
        cell.text = str(header)
        _set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(
                    int(fg[0:2], 16),
                    int(fg[2:4], 16),
                    int(fg[4:6], 16)
                )


def _safe_cell(cell, value):
    """Safely set cell text — always converts to string."""
    cell.text = str(value) if value is not None else "—"


def generate_booklet(eng, tasks, comments_by_task, reviews_by_task,
                     queries, team, filepath,
                     firm_name="", firm_reg_no=""):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AUDIT WORKING PAPERS BOOKLET")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(str(eng.get("client_name") or ""))
    r.bold = True
    r.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run(
        f"{eng.get('engagement_type', '')} — FY {eng.get('financial_year', '')}"
    )
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    # Firm info table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    tbl.rows[0].cells[0].text = "Prepared by"
    tbl.rows[0].cells[1].text = str(firm_name) if firm_name else "CA Firm"
    if firm_reg_no:
        r_ = tbl.add_row()
        r_.cells[0].text = "Firm Reg. No."
        r_.cells[1].text = str(firm_reg_no)
    r2_ = tbl.add_row()
    r2_.cells[0].text = "Generated on"
    r2_.cells[1].text = datetime.now().strftime("%d %B %Y, %I:%M %p")
    r3_ = tbl.add_row()
    r3_.cells[0].text = "Period"
    r3_.cells[1].text = (
        f"{eng.get('period_from', '—')} to {eng.get('period_to', '—')}"
    )

    doc.add_page_break()

    # ── SECTION 1: ENGAGEMENT DETAILS ───────────────────────
    _heading(doc, "1. Engagement Details")
    _kv(doc, "Client",          eng.get("client_name", ""))
    _kv(doc, "PAN",             eng.get("pan", ""))
    _kv(doc, "GSTIN",           eng.get("gstin", ""))
    _kv(doc, "Engagement Type", eng.get("engagement_type", ""))
    _kv(doc, "Financial Year",  eng.get("financial_year", ""))
    _kv(doc, "Period",
        f"{eng.get('period_from', '—')} to {eng.get('period_to', '—')}")
    _kv(doc, "Status",          eng.get("status", ""))
    if eng.get("notes"):
        _kv(doc, "Notes", eng.get("notes", ""))

    doc.add_paragraph()

    # ── SECTION 2: AUDIT TEAM ────────────────────────────────
    _heading(doc, "2. Audit Team")
    if team:
        tbl2 = doc.add_table(rows=1, cols=3)
        tbl2.style = 'Table Grid'
        _tbl_header(tbl2, ["Name", "Role", "Email"])
        for m in team:
            row_ = tbl2.add_row()
            _safe_cell(row_.cells[0], m.get("full_name", ""))
            _safe_cell(row_.cells[1], m.get("role", ""))
            _safe_cell(row_.cells[2], m.get("email", ""))
    else:
        doc.add_paragraph("No team members assigned.")

    doc.add_page_break()

    # ── SECTION 3: TASK WORKING PAPERS ──────────────────────
    _heading(doc, "3. Task-wise Working Papers")

    # Group tasks by area
    areas = {}
    for t in tasks:
        area = t.get("area") or "General"
        areas.setdefault(area, []).append(t)

    for area, area_tasks in areas.items():
        _heading(doc, area, level=2, color="1a4d80")

        for t in area_tasks:
            tid = t["id"]

            # Task title
            p = doc.add_paragraph()
            r_ = p.add_run(f"  {t.get('title', '')}")
            r_.bold = True
            r_.font.size = Pt(10)
            r_.font.color.rgb = RGBColor(0, 51, 102)

            # Task detail table
            dtbl = doc.add_table(rows=1, cols=4)
            dtbl.style = 'Table Grid'
            _tbl_header(dtbl,
                        ["Status", "Priority", "Due Date", "WP Ref"],
                        bg="1a4d80", fg="DAA520")
            drow = dtbl.add_row()
            _safe_cell(drow.cells[0], t.get("status", ""))
            _safe_cell(drow.cells[1], t.get("priority", ""))
            _safe_cell(drow.cells[2], t.get("due_date") or "—")
            _safe_cell(drow.cells[3], t.get("working_paper_ref") or "—")

            # Assignees
            _kv(doc, "    Assigned To",
                t.get("assignee_name", "") or "Unassigned")

            if t.get("description"):
                _kv(doc, "    Description", t["description"])

            # Comments
            cmts = comments_by_task.get(tid, [])
            if cmts:
                cp = doc.add_paragraph("    Comments:")
                cp.runs[0].bold = True
                for c in cmts:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Cm(1)
                    r1 = p2.add_run(
                        f"[{c.get('author_name', '?')} — "
                        f"{str(c.get('created_at', ''))[:16]}]: "
                    )
                    r1.bold = True
                    r1.font.size = Pt(9)
                    r1.font.color.rgb = RGBColor(80, 80, 80)
                    p2.add_run(str(c.get("content", "")))

            # Reviews
            revs = reviews_by_task.get(tid, [])
            if revs:
                rp = doc.add_paragraph("    Reviews:")
                rp.runs[0].bold = True
                for rv in revs:
                    p3 = doc.add_paragraph()
                    p3.paragraph_format.left_indent = Cm(1)
                    action = str(rv.get("action", ""))
                    color = (RGBColor(30, 126, 68)
                             if action == "Approved"
                             else RGBColor(197, 48, 48))
                    r1 = p3.add_run(f"[{action}] ")
                    r1.bold = True
                    r1.font.color.rgb = color
                    r1.font.size = Pt(9)
                    p3.add_run(
                        f"{rv.get('reviewer_name', '')} — "
                        f"{str(rv.get('reviewed_at', ''))[:16]}"
                    )
                    if rv.get("remarks"):
                        p3.add_run(f" — {rv['remarks']}")

            doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 4: QUERY SHEET ───────────────────────────────
    _heading(doc, "4. Query Sheet")
    if queries:
        qtbl = doc.add_table(rows=1, cols=6)
        qtbl.style = 'Table Grid'
        _tbl_header(qtbl,
                    ["Sr.", "Query", "Response",
                     "Status", "Raised By", "Date"])
        for q in queries:
            qrow = qtbl.add_row()
            _safe_cell(qrow.cells[0], q.get("sr_no") or "")
            _safe_cell(qrow.cells[1], q.get("query_text") or "")
            _safe_cell(qrow.cells[2], q.get("response") or "Pending")
            _safe_cell(qrow.cells[3], q.get("status") or "")
            _safe_cell(qrow.cells[4], q.get("raised_by_name") or "")
            _safe_cell(qrow.cells[5], str(q.get("raised_date") or "")[:10])
    else:
        doc.add_paragraph("No queries raised for this engagement.")

    # ── FOOTER ──────────────────────────────────────────────
    doc.add_page_break()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Generated by CA FirmHub"
        + (f" | {firm_name}" if firm_name else "")
        + f" | {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(150, 150, 150)

    doc.save(filepath)
